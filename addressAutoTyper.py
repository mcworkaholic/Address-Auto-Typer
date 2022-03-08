import pandas as pd
import keyboard
from docx import Document
from docx.shared import Pt
from termcolor import colored
import os
from openpyxl import Workbook


keyboard.press('F11')
line_break = '\n'
os.system('color')
string_one = "Address Auto Typer found the following spreadsheet files..."
string_two = "Which file would you like to prepare? (insert # only) "
string_four = "Press enter to continue."

print('\n')
print(string_one.center(150), line_break)

fileList = []
for root, dirs, files in os.walk(r'C:\Users\wevan\OneDrive'):
    for file in files:
        if file.endswith(".xlsx") or file.endswith(".GSHEET"):
            fileList.append(file)
for i, value in enumerate(fileList):
    print(i, "-", fileList[i])
print('\n')

fileSelection = int(input(string_two.center(55)))
fileName = fileList[fileSelection]
table = pd.read_excel("C:\\Users\wevan\OneDrive\Desktop\\" + fileName)
addressList = table["address"].tolist()
longest_string = max(addressList, key=len)
lsLength = len(longest_string)

def print_test():
    print('\n')
    print(table, line_break)
    print(addressList, line_break)
    print("address one = " + addressList[0], line_break)
    print("last address = " + addressList[-1], line_break)
    print("longest string = " + str(longest_string), line_break)
    print("longest string char length = " + str(lsLength))

def convert(string):
    fileList2 = []
    orderNum = []
    fileList2[:0] = string
    for i, value in enumerate(fileName):
        if fileList2[i].isdigit():
            orderNum.append(fileList2[i])
    global order
    order = (''.join(str(x) for x in orderNum))
    return order

def excel_format():
    wb = Workbook()
    ws = wb.create_sheet("Check-List", 0)
    ws['A1'] = 'Address'
    ws['B1'] = 'Before'
    ws['C1'] = 'During'
    ws['D1'] = 'After'
    ws['E1'] = 'COA'
    ws['F1'] = 'Comments'
    ws.column_dimensions['A'].width = lsLength + 3
    ws.column_dimensions['F'].width = 25
    for i, value in enumerate(addressList, start=2):
     ws.cell(row=i, column=1).value = value
    wb.save('C:\\Users\wevan\OneDrive\Desktop\\Workorderchecklist.xlsx')

def document_format():
    count = 0
    document = Document()
    styleEventStamp = document.styles['Normal']
    font = styleEventStamp.font
    font.name = "Bahnschrift"
    font.size = Pt(40)

    for i, value in enumerate(addressList):
        count += 1
        eventStampBefore = document.add_paragraph('Before             wo #' + order)
        address = document.add_paragraph(addressList[i])
        address.style = document.styles['Normal']
        eventStampBefore.styleEventStamp = document.styles['Normal']
        run = address.add_run()
        run.add_break()
        document.add_paragraph('During             wo #' + order)
        address = document.add_paragraph(addressList[i])
        run = address.add_run()
        run.add_break()
        document.add_paragraph('After                wo #' + order)
        document.add_paragraph(addressList[i])
        document.add_page_break()
    document.save('C:\\Users\wevan\OneDrive\Desktop\\Workorder.docx')

def driver():
    print_test()
    convert(fileName)
    document_format()
    excel_format()
    finish()

def finish():
    print('\n')
    input(colored(string_four.center(150), 'green'))
    os.system('start C:\\Users\wevan\OneDrive\Desktop\\Workorder.docx')
    os.system('start C:\\Users\wevan\OneDrive\Desktop\\Workorderchecklist.xlsx')
    os.system("TASKKILL /F /IM AddressTyper.exe")

driver()
